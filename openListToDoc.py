import xlrd
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor

xlsx = xlrd.open_workbook('D:\\GitBase\\ArticleCheck\\list.xls')
sheet1 = xlsx.sheets()[0]
sheet1_nrows = sheet1.nrows

doc1 = Document()
for i in range(sheet1_nrows):  # 逐行打印sheet1数据
    
    if i <=234 and i >=198:
        doc1.styles['Normal'].font.name = u'宋体'
        doc1.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        doc1.styles['Normal'].font.size = Pt(12)

        doc1.add_paragraph('项目编号: %s'%sheet1.row_values(i)[1])
        doc1.add_paragraph('项目名称: %s'%sheet1.row_values(i)[2])
        doc1.add_paragraph('所属行业: %s'%sheet1.row_values(i)[3])
        doc1.add_paragraph('项目概述:')
        doc1.add_paragraph(sheet1.row_values(i)[4])

        doc1.add_paragraph(' ')
        doc1.add_paragraph(' ')

        filename = sheet1.row_values(i)[0]
        filename = ''.join(filename)

doc1.save('./article/taoyy.docx')
#sys.exit()