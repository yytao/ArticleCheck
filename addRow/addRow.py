from docx import Document


document = Document('./result.docx')

for paragraph in document.paragraphs:
    if "所属行业" in paragraph.text:
        t1 = paragraph.insert_paragraph_before("")
        t2 = paragraph.insert_paragraph_before("")
        
        t1.add_run('第一参赛者姓名：').bold = True
        t2.add_run('现所在国家/地区：').bold = True
        
document.save('./new/result.docx')


